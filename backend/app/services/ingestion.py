import uuid
import os
import io
from io import BytesIO
from fastapi import UploadFile
from sqlalchemy.orm import Session
from ..infrastructure.database import Document, IngestionTask
# Import moved inside function to avoid circular import
from pdfminer.high_level import extract_text as pdf_extract_text
import docx
from PIL import Image
import pytesseract



def extract_text(file: UploadFile) -> str:
    """Extract plain text from supported file types.
    Supports PDF, DOCX, and image files (OCR). Returns empty string on failure.
    """
    suffix = os.path.splitext(file.filename)[1].lower()
    # Read raw bytes once
    raw = file.file.read()
    # Reset file pointer for later use if needed
    file.file.seek(0)
    if suffix == ".pdf":
        # pdfminer works with file path, so write to temp
        tmp_path = os.path.join(os.getenv("TMPDIR", "/tmp"), f"{uuid.uuid4()}.pdf")
        with open(tmp_path, "wb") as f:
            f.write(raw)
        try:
            return pdf_extract_text(tmp_path) or ""
        finally:
            os.remove(tmp_path)
    if suffix in {".docx", ".doc"}:
        tmp_path = os.path.join(os.getenv("TMPDIR", "/tmp"), f"{uuid.uuid4()}.docx")
        with open(tmp_path, "wb") as f:
            f.write(raw)
        try:
            doc = docx.Document(tmp_path)
            return "\n".join(p.text for p in doc.paragraphs)
        finally:
            os.remove(tmp_path)
    if suffix in {".png", ".jpg", ".jpeg", ".tiff", ".bmp"}:
        img = Image.open(BytesIO(raw))
        return pytesseract.image_to_string(img)
    # Fallback: try UTF‑8 decode
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError:
        return ""


def enqueue_ingestion(file: UploadFile, db: Session) -> str:
    """Persist uploaded file metadata, create IngestionTask, and dispatch a Celery job.
    Returns the task UUID.
    """
    uploads_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../uploads"))
    os.makedirs(uploads_dir, exist_ok=True)
    file_path = os.path.join(uploads_dir, file.filename)
    with open(file_path, "wb") as out_f:
        out_f.write(file.file.read())

    # Create Document record
    doc = Document(
        filename=file.filename,
        content_type=file.content_type or "application/octet-stream",
        path=file_path,
    )
    db.add(doc)
    db.flush()  # get doc.id

    task_id = str(uuid.uuid4())
    ingestion_task = IngestionTask(
        id=task_id,
        document_id=doc.id,
        status="queued",
    )
    db.add(ingestion_task)
    db.commit()

    # Dispatch Celery job
    from ..workers.celery_app import celery_app
    celery_app.send_task("tasks.process_document", args=[task_id])
    return task_id
